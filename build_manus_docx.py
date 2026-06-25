#!/usr/bin/env python3
"""Bygger manus.docx (presentasjonsark) fra manus.md.

Fargekoding:
  - SORT  = det du faktisk skal SI (vanlig tekst, «sitater», **uthevet**).
  - RØD   = det du skal HUSKE (regi i *kursiv*, [hakeparentes], HUSK-lapper,
            overskrifter, kodeblokker).

Layout:
  - LIGGENDE (landskap) A4.
  - Seksjoner pakkes sammen på arkene (mindre blaing) – IKKE ett ark per seksjon.
  - Mellom hver seksjon: en stiplet ✂ KLIPPELINJE, så du kan klippe arket i
    biter. keepNext/keepLines holder hver seksjon samlet, så brudd skjer ved
    klippelinjene – ikke midt i en seksjon.
Kjør: python3 build_manus_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "manus.md"
OUT = "manus.docx"

RED = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)

SAY_SIZE = 12.5    # det du sier
REGI_SIZE = 10     # det du husker
MONO_SIZE = 9

INLINE = re.compile(r'\*\*(.+?)\*\*|\*([^*]+?)\*')


def style_say(run, bold=False):
    run.font.color.rgb = BLACK
    run.font.size = Pt(SAY_SIZE)
    run.bold = bold


def style_regi(run):
    run.font.color.rgb = RED
    run.font.size = Pt(REGI_SIZE)
    run.italic = True


def add_inline(p, text):
    """Legg til tekst-runs: *kursiv* -> rød (husk), **fet**/vanlig -> sort (si)."""
    text = text.replace('`', '').replace('[x]', '✔').replace('[ ]', '☐')
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            style_say(p.add_run(text[pos:m.start()]))
        if m.group(1) is not None:        # **fet** -> sort uthevet
            style_say(p.add_run(m.group(1)), bold=True)
        else:                              # *kursiv* -> rød regi
            style_regi(p.add_run(m.group(2)))
        pos = m.end()
    if pos < len(text):
        style_say(p.add_run(text[pos:]))


def keep(p, with_next=True):
    """Hold avsnittet samlet (keepLines) og evt. sammen med neste (keepNext)."""
    pPr = p._p.get_or_add_pPr()
    if with_next:
        pPr.append(OxmlElement('w:keepNext'))
    pPr.append(OxmlElement('w:keepLines'))
    return p


def add_cut_line(doc):
    """Stiplet ✂ klippelinje mellom seksjoner. Brudd kan skje her (ikke keepNext)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run('✂ ' + ' ' * 4)
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'dashed')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '999999')
    pbdr.append(bottom)
    pPr.append(pbdr)
    keep(p, with_next=False)
    return p


def _add_field(paragraph, instr):
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr_el = OxmlElement('w:instrText')
    instr_el.set(qn('xml:space'), 'preserve'); instr_el.text = instr
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r.append(begin); r.append(instr_el); r.append(end)
    return run


def add_page_numbers(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = [p.add_run('Side '), _add_field(p, 'PAGE'),
             p.add_run(' / '), _add_field(p, 'NUMPAGES')]
    for run in parts:
        run.font.size = Pt(9)
        run.font.color.rgb = GREY


def main():
    with open(SRC, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()

    # LIGGENDE A4 + margfeatures
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Mm(297)
    sec.page_height = Mm(210)
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Mm(14))
    add_page_numbers(sec)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(SAY_SIZE)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    i = 0
    card_no = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Kodeblokk ```
        if stripped.startswith('```'):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            for cl in buf:
                p = doc.add_paragraph()
                r = p.add_run(cl if cl else ' ')
                r.font.name = 'Consolas'
                r.font.size = Pt(MONO_SIZE)
                r.font.color.rgb = RED
                keep(p)
            continue

        # Overskrift
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped[level:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            if level == 2:                       # kort-nummer pr. seksjon
                card_no += 1
                badge = p.add_run(f'[{card_no}]  ')
                badge.bold = True
                badge.font.color.rgb = GREY
                badge.font.size = Pt(13)
            r = p.add_run(text.replace('`', ''))
            r.bold = True
            r.font.color.rgb = RED
            r.font.size = Pt({1: 16, 2: 13}.get(level, 12))
            keep(p)             # overskrift følger med innholdet under
            i += 1
            continue

        # Skillelinje --- -> klippelinje
        if stripped == '---':
            add_cut_line(doc)
            i += 1
            continue

        if stripped == '':
            i += 1
            continue

        # Blockquote (eksempel-/sitatlinjer)
        if stripped.startswith('>'):
            buf = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                buf.append(lines[i].strip().lstrip('>').strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(5)
            add_inline(p, ' '.join(buf))
            keep(p)
            continue

        # Listepunkt
        if re.match(r'^\s*[-*]\s+', line):
            content = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(4)
            style_say(p.add_run('• '))
            add_inline(p, content)
            keep(p)
            i += 1
            continue

        # Vanlig avsnitt: slå sammen fortsettelseslinjer
        buf = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            ns = nxt.strip()
            if ns == '' or ns.startswith('#') or ns == '---' \
                    or ns.startswith('>') or ns.startswith('```') \
                    or re.match(r'^\s*[-*]\s+', nxt):
                break
            buf.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, ' '.join(s.strip() for s in buf))
        keep(p)

    doc.save(OUT)
    print(f"Skrev {OUT}")


if __name__ == '__main__':
    main()
